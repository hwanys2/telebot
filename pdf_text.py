import pdftotext
import os
import docx
path = "/Users/jinhwan/Desktop/PISA 2021 Mathematics Framework Draft.pdf"
folder = "/".join(path.split("/")[0:-1]) + "/"
file_name = path.split("/")[-1].split(".")[0]
print(file_name)

with open(path, "rb") as f:
    pdf = pdftotext.PDF(f)

pdf = "\n\n".join(pdf)
pdf.rstrip('\n')
print(pdf)
with open(folder + file_name + ".txt", 'w') as f:
    f.write(pdf)

#docx로 저장하기
# doc = docx.Document()
# para = doc.add_paragraph()
# run = para.add_run(pdf)
# doc.save('example.docx')